"""Accepted-row CSV/XLSX and mailing-label DOCX export."""

from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from . import statuses
from .quality import CONTACT_PREFIXES, clean_text, contact_key, mailable_contact, raw_text_issue
from .storage import ReviewStore


EXPORT_FIELDS = [
    "meeting_date",
    "event_id",
    "city_item_id",
    "file_id",
    "agenda_description",
    "applicant_name",
    "applicant_company",
    "applicant_mailing_address",
    "applicant_phone",
    "applicant_email",
    "project_contact_name",
    "project_contact_company",
    "project_contact_mailing_address",
    "project_contact_phone",
    "project_contact_email",
    "owner_name",
    "owner_company",
    "owner_mailing_address",
    "owner_phone",
    "owner_email",
    "section5_description",
    "unit_count",
    "source_url",
    "notes",
]

@dataclass(frozen=True)
class LabelContact:
    """Purpose: carry one cleaned contact destined for an address label."""

    extraction_id: int
    contact_type: str
    name: str
    company: str
    address: str

    @property
    def key(self) -> str:
        """Purpose: dedupe labels after whitespace and case normalization."""

        return contact_key(self.name, self.company, self.address)

    def lines(self) -> list[str]:
        """Purpose: format label lines for Avery address cells."""

        return [line for line in (self.name, self.company, self.address) if line]


class ExportService:
    """Purpose: export reviewed workbench data without runtime dependencies."""

    def __init__(self, store: ReviewStore) -> None:
        self.store = store

    def export(self, output_path: Path, status: str = statuses.ACCEPTED) -> dict[str, Any]:
        """Purpose: write accepted rows to CSV or XLSX and record metadata."""

        rows = self.store.accepted_export_rows(status)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        fmt = output_path.suffix.lower().lstrip(".") or "csv"
        if fmt == "csv":
            self._write_csv(output_path, rows)
            result_extra: dict[str, Any] = {}
        elif fmt == "xlsx":
            self._write_xlsx(output_path, rows)
            result_extra = {}
        elif fmt == "docx":
            contacts, issues = self._label_contacts(rows)
            if not contacts:
                raise ValueError("No accepted contacts passed mailing-label quality control")
            self._write_avery_5160_docx(output_path, contacts)
            rows = [contact.__dict__ for contact in contacts]
            result_extra = {"qc_skipped_count": len(issues), "qc_issues": issues, "label_template": "avery_5160"}
        else:
            raise ValueError("Export output must end in .csv, .xlsx, or .docx")
        export_id = self.store.record_export(output_path, fmt, len(rows))
        return {"id": export_id, "path": str(output_path), "format": fmt, "row_count": len(rows), **result_extra}

    def _write_csv(self, output_path: Path, rows: list[dict[str, Any]]) -> None:
        """Purpose: write a portable comma-separated export."""

        with output_path.open("w", newline="", encoding="utf-8") as fh:
            writer = csv.DictWriter(fh, fieldnames=EXPORT_FIELDS, extrasaction="ignore")
            writer.writeheader()
            for row in rows:
                writer.writerow(row)

    def _write_xlsx(self, output_path: Path, rows: list[dict[str, Any]]) -> None:
        """Purpose: write an Excel workbook for DBMS import."""

        try:
            from openpyxl import Workbook
        except Exception as exc:
            raise RuntimeError("openpyxl is required for .xlsx export") from exc
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "accepted"
        sheet.append(EXPORT_FIELDS)
        for row in rows:
            sheet.append([row.get(field) for field in EXPORT_FIELDS])
        workbook.save(output_path)

    def _label_contacts(self, rows: list[dict[str, Any]]) -> tuple[list[LabelContact], list[dict[str, Any]]]:
        """Purpose: keep only accepted contacts that are clean enough to mail."""

        contacts: list[LabelContact] = []
        issues: list[dict[str, Any]] = []
        seen: set[str] = set()
        for row in self._newest_rows_first(rows):
            for prefix in CONTACT_PREFIXES:
                contact, reason = self._label_contact(row, prefix)
                if not contact:
                    if reason:
                        issues.append({"extraction_id": row.get("id"), "contact_type": prefix, "reason": reason})
                    continue
                if contact.key in seen:
                    issues.append(
                        {
                            "extraction_id": row.get("id"),
                            "contact_type": prefix,
                            "reason": "outdated duplicate contact; newer accepted contact retained",
                        }
                    )
                    continue
                seen.add(contact.key)
                contacts.append(contact)
        return contacts, issues

    def _newest_rows_first(self, rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
        """Purpose: prefer recent accepted contact data during label dedupe."""

        return sorted(rows, key=lambda row: (str(row.get("meeting_date") or ""), int(row.get("id") or 0)), reverse=True)

    def _label_contact(self, row: dict[str, Any], prefix: str) -> tuple[LabelContact | None, str | None]:
        """Purpose: validate one applicant/project-contact/owner mailing record."""

        contact = mailable_contact(row, prefix)
        name = clean_text(row.get(f"{prefix}_name"))
        company = clean_text(row.get(f"{prefix}_company"))
        address = clean_text(row.get(f"{prefix}_mailing_address"))
        if not any((name, company, address)):
            return None, None
        if not contact:
            return None, "missing address or name/company"
        raw_issue = raw_text_issue(name, company, address)
        if raw_issue:
            return None, raw_issue
        return LabelContact(int(row["id"]), prefix, name, company, address), None

    def _write_avery_5160_docx(self, output_path: Path, contacts: list[LabelContact]) -> None:
        """Purpose: write Avery 5160/8160-style 30-up address labels."""

        try:
            from docx import Document
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt
        except Exception as exc:
            raise RuntimeError("python-docx is required for .docx label export") from exc

        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.1875)
        section.right_margin = Inches(0.1875)
        table = document.add_table(rows=10, cols=5)
        table.autofit = False
        self._set_table_borders(table, OxmlElement, qn)
        widths = (2.625, 0.125, 2.625, 0.125, 2.625)
        for row in table.rows:
            row.height = Inches(1)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for idx, cell in enumerate(row.cells):
                cell.width = Inches(widths[idx])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(cell, OxmlElement, qn)
        label_cells = [cell for row in table.rows for idx, cell in enumerate(row.cells) if idx in {0, 2, 4}]
        for cell, contact in zip(label_cells, contacts):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            for idx, line in enumerate(contact.lines()):
                if idx:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(line)
                run.font.name = "Arial"
                run.font.size = Pt(10)
        document.save(output_path)

    def _set_table_borders(self, table: Any, element: Any, qname: Any) -> None:
        """Purpose: remove table lines so Word prints label text only."""

        borders = element("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = element(f"w:{side}")
            border.set(qname("w:val"), "nil")
            borders.append(border)
        table._tbl.tblPr.append(borders)

    def _set_cell_margins(self, cell: Any, element: Any, qname: Any) -> None:
        """Purpose: keep text away from physical label edges."""

        tc_pr = cell._tc.get_or_add_tcPr()
        margins = element("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            margin = element(f"w:{side}")
            margin.set(qname("w:w"), "72")
            margin.set(qname("w:type"), "dxa")
            margins.append(margin)
        tc_pr.append(margins)
