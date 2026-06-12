from __future__ import annotations

import datetime as dt
import json
import sqlite3
from pathlib import Path
import zipfile

from plan_commission_workbench import statuses
from plan_commission_workbench.api import PlanCommissionWorkbench
from plan_commission_workbench.export import ExportService
from plan_commission_workbench.models import (
    AgendaClassification,
    AgendaSegment,
    ApplicationExtraction,
    ContactFields,
    FieldEvidence,
    RunRequest,
)
from plan_commission_workbench.storage import ReviewStore, _pid_alive
from plan_commission_workbench.runtime import WorkbenchRuntime
from plan_commission_workbench.watchdog import RunWatchdog


def test_store_dedupe_review_and_export(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), None)
    source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id="27999",
        file_id=None,
        attachment_id=None,
        source_url="https://example.test/agenda.pdf",
        content_hash="agenda-hash",
        processing_status=statuses.AGENDA_HIT,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        source_id,
        AgendaSegment("27999", "96005", "88001", dt.date(2026, 6, 1), "Construct apartments"),
        AgendaClassification("96005", statuses.AGENDA_HIT, 0.91, "Housing", "Construct apartments"),
    )

    assert store.agenda_complete("27999", source_url="https://example.test/agenda.pdf")

    app_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="27999",
        file_id="88001",
        attachment_id="171817",
        source_url="https://example.test/application.pdf",
        content_hash="app-hash",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    extraction_id = store.upsert_application_extraction(
        run_id,
        app_source,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url="https://example.test/application.pdf",
            attachment_id="171817",
            applicant=ContactFields(
                name="Jane Applicant",
                company="Applicant LLC",
                mailing_address="123 Main Street, Madison, WI 53703",
            ),
            project_contact=ContactFields(email="pat@example.com"),
            owner=ContactFields(),
            section5_description="Construct 48 dwelling units.",
            unit_count=48,
            status=statuses.APPLICATION_EXTRACTED,
            target_project=True,
            evidence=(FieldEvidence("unit_count", 48, "48 dwelling units", 0.9),),
        ),
    )

    assert store.application_complete(agenda_id, "https://example.test/application.pdf", "171817")

    store.review_application(
        extraction_id,
        statuses.ACCEPTED,
        {"applicant_name": "Jane Corrected"},
        "Looks clean",
    )
    export_path = tmp_path / "accepted.csv"
    result = ExportService(store).export(export_path, statuses.ACCEPTED)

    assert result["row_count"] == 1
    body = export_path.read_text(encoding="utf-8")
    assert "Jane Corrected" in body
    assert "Construct 48 dwelling units" in body


def test_label_export_uses_corrected_clean_contacts_and_reports_qc(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), None)
    agenda_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id="27999",
        file_id=None,
        attachment_id=None,
        source_url="https://example.test/agenda.pdf",
        content_hash="agenda-hash",
        processing_status=statuses.AGENDA_HIT,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        agenda_source,
        AgendaSegment("27999", "96005", "88001", dt.date(2026, 6, 1), "Construct apartments"),
        AgendaClassification("96005", statuses.AGENDA_HIT, 0.91, "Housing", "Construct apartments"),
    )
    app_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="27999",
        file_id="88001",
        attachment_id="171817",
        source_url="https://example.test/application.pdf",
        content_hash="app-hash",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    extraction_id = store.upsert_application_extraction(
        run_id,
        app_source,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url="https://example.test/application.pdf",
            attachment_id="171817",
            applicant=ContactFields(name="Applicant name Jane Raw"),
            project_contact=ContactFields(name="Pat Contact", email="pat@example.com"),
            owner=ContactFields(),
            section5_description="Construct 48 dwelling units.",
            unit_count=48,
            status=statuses.APPLICATION_EXTRACTED,
            target_project=True,
        ),
    )
    store.review_application(
        extraction_id,
        statuses.ACCEPTED,
        {
            "applicant_name": "Jane Corrected",
            "applicant_company": "Clean Housing LLC",
            "applicant_mailing_address": "123 Main Street, Madison, Wisconsin",
        },
        None,
    )

    export_path = tmp_path / "labels.docx"
    result = ExportService(store).export(export_path, statuses.ACCEPTED)

    assert result["format"] == "docx"
    assert result["row_count"] == 1
    assert result["qc_skipped_count"] == 1
    assert result["qc_issues"][0]["contact_type"] == "project_contact"
    body = _docx_text(export_path)
    assert "Jane Corrected" in body
    assert "Clean Housing LLC" in body
    assert "123 Main Street, Madison, Wisconsin" in body
    assert "Pat Contact" not in body


def test_label_export_allows_company_only_contacts(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), None)
    _accepted_extraction(
        store,
        run_id,
        meeting_date=dt.date(2026, 6, 1),
        city_item_id="96005",
        applicant=ContactFields(
            company="Company Only Housing LLC",
            mailing_address="123 Main Street, Madison, Wisconsin",
        ),
        project_contact=ContactFields(),
    )

    result = ExportService(store).export(tmp_path / "labels.docx", statuses.ACCEPTED)
    body = _docx_text(tmp_path / "labels.docx")

    assert result["row_count"] == 1
    assert result["qc_issues"] == []
    assert "Company Only Housing LLC" in body
    assert "123 Main Street, Madison, Wisconsin" in body


def test_data_relative_export_path_uses_runtime_data_dir(tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)

    path = workbench._export_path(Path("data/exports/madison_review.xlsx"))

    assert path == tmp_path / "user-data" / "exports" / "madison_review.xlsx"


def test_diagnostic_bundle_contains_restorable_db_and_logs(tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)
    workbench.store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), "debug")
    runtime.server_log_path.write_text("server log", encoding="utf-8")
    runtime.server_error_log_path.write_text("server error", encoding="utf-8")

    result = workbench.create_diagnostic_bundle()

    bundle_path = Path(result["path"])
    assert bundle_path.exists()
    with zipfile.ZipFile(bundle_path) as archive:
        names = set(archive.namelist())
        assert {"workbench.db", "manifest.json", "server.log", "server.err.log"} <= names
        manifest = json.loads(archive.read("manifest.json"))
        assert manifest["latest_runs"][0]["run_request_text"] == "debug"
        extracted_db = tmp_path / "restored.db"
        extracted_db.write_bytes(archive.read("workbench.db"))
    with sqlite3.connect(extracted_db) as conn:
        assert conn.execute("SELECT COUNT(*) FROM runs").fetchone()[0] == 1


def test_diagnostic_bundle_survives_locked_log_file(monkeypatch, tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)
    runtime.server_log_path.write_text("server log", encoding="utf-8")
    original_read_text = Path.read_text

    def guarded_read_text(path: Path, *args, **kwargs):
        if path == runtime.server_log_path:
            raise OSError("file is locked")
        return original_read_text(path, *args, **kwargs)

    monkeypatch.setattr(Path, "read_text", guarded_read_text)

    result = workbench.create_diagnostic_bundle()

    with zipfile.ZipFile(result["path"]) as archive:
        assert "server.log.error.txt" in archive.namelist()
        assert "file is locked" in archive.read("server.log.error.txt").decode("utf-8")


def test_diagnostic_bundle_downloads_with_backup_warning(monkeypatch, tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)

    def fail_backup(_destination: Path) -> Path:
        raise OSError("database is locked")

    monkeypatch.setattr(workbench.store, "backup_to", fail_backup)

    result = workbench.create_diagnostic_bundle()

    assert "database is locked" in result["warning"]
    with zipfile.ZipFile(result["path"]) as archive:
        assert "workbench.db" not in archive.namelist()
        manifest = json.loads(archive.read("manifest.json"))
        assert manifest["db_included"] is False
        assert "database is locked" in manifest["backup_error"]


def test_diagnostic_bundle_downloads_when_temp_cleanup_fails(monkeypatch, tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)

    monkeypatch.setattr(
        workbench,
        "_remove_temp_bundle_file",
        lambda path: f"Temporary backup cleanup failed for {path.name}: locked",
    )

    result = workbench.create_diagnostic_bundle()

    assert "Temporary backup cleanup failed" in result["warning"]
    with zipfile.ZipFile(result["path"]) as archive:
        assert "workbench.db" in archive.namelist()


def test_diagnostic_bundle_includes_run_worker_logs(tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)
    stdout_path, stderr_path = runtime.run_worker_log_paths(42)
    stdout_path.write_text("worker stdout", encoding="utf-8")
    stderr_path.write_text("worker stderr", encoding="utf-8")

    result = workbench.create_diagnostic_bundle()

    with zipfile.ZipFile(result["path"]) as archive:
        assert archive.read("run_logs/run_42.log").decode("utf-8") == "worker stdout"
        assert archive.read("run_logs/run_42.err.log").decode("utf-8") == "worker stderr"


def test_start_madison_run_worker_records_child_pid(monkeypatch, tmp_path) -> None:
    runtime = WorkbenchRuntime(project_root=tmp_path / "bundle", data_dir=tmp_path / "user-data")
    workbench = PlanCommissionWorkbench(runtime=runtime)
    run_id = workbench.create_madison_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2))
    launched = {}

    class FakePopen:
        """Purpose: capture worker launch arguments without starting a scrape."""

        pid = 2468

        def __init__(self, command, stdout, stderr, env, text, **kwargs) -> None:
            launched["command"] = command
            launched["stdout"] = Path(stdout.name).name
            launched["stderr"] = Path(stderr.name).name
            launched["env"] = env
            launched["text"] = text
            launched["process_kwargs"] = kwargs

    monkeypatch.setattr("plan_commission_workbench.api.subprocess.Popen", FakePopen)

    result = workbench.start_madison_run_worker(
        run_id,
        RunRequest(dt.date(2026, 6, 1), dt.date(2026, 6, 2)),
    )

    row = workbench.store.get_run(run_id)
    assert result["worker_pid"] == 2468
    assert row and row["worker_pid"] == 2468
    assert launched["command"][1:3] == ["-m", "plan_commission_workbench.run_worker"]
    assert launched["env"]["PCW_DATA_DIR"] == str(runtime.data_dir)
    assert launched["stdout"] == f"run_{run_id}.log"
    assert launched["stderr"] == f"run_{run_id}.err.log"
    assert launched["text"] is True
    assert "start_new_session" in launched["process_kwargs"] or "creationflags" in launched["process_kwargs"]


def test_store_backup_closes_destination_connection(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    backup_path = tmp_path / "backup.db"

    store.backup_to(backup_path)
    backup_path.unlink()

    assert not backup_path.exists()


def test_application_sources_allow_duplicate_content_hashes(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 1, 1), dt.date(2026, 1, 31), None)
    first = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="1",
        file_id="100",
        attachment_id="a",
        source_url="https://example.test/application-a.pdf",
        content_hash="same-application-hash",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    second = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="2",
        file_id="200",
        attachment_id="b",
        source_url="https://example.test/application-b.pdf",
        content_hash=None,
        processing_status=statuses.APPLICATION_DOWNLOADING,
    )

    updated = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="2",
        file_id="200",
        attachment_id="b",
        source_url="https://example.test/application-b.pdf",
        content_hash="same-application-hash",
        processing_status=statuses.APPLICATION_DOCLING,
    )

    assert first != second
    assert updated == second
    with sqlite3.connect(store.db_path) as conn:
        rows = conn.execute(
            "SELECT id, source_url FROM source_items WHERE source_kind = 'application' AND content_hash = ? ORDER BY id",
            ("same-application-hash",),
        ).fetchall()
    assert [row[0] for row in rows] == [first, second]


def test_startup_migration_downgrades_public_comment_hits_and_counters(tmp_path) -> None:
    db_path = tmp_path / "workbench.db"
    store = ReviewStore(db_path)
    store.initialize()
    run_id = store.create_run(dt.date(2025, 12, 1), dt.date(2025, 12, 1), None)
    source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id="27921",
        file_id=None,
        attachment_id=None,
        source_url="https://example.test/agenda.pdf",
        content_hash="agenda-hash",
        processing_status=statuses.AGENDA_HIT,
    )
    store.upsert_agenda_item(
        run_id,
        source_id,
        AgendaSegment("27921", "71173", "60306", dt.date(2025, 12, 1), "Plan Commission Public Comment Period"),
        AgendaClassification("71173", statuses.AGENDA_HIT, 0.8, "Public comment", "Public comment"),
    )
    store.update_counters(run_id)

    reopened = ReviewStore(db_path)
    reopened.initialize()

    run = reopened.get_run(run_id)
    assert run and run["agenda_hits"] == 0
    assert reopened.list_agenda_items(statuses.NOT_TARGET_PROJECT)[0]["city_item_id"] == "71173"


def test_agenda_review_approval_unsticks_related_clean_application(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 5, 11), dt.date(2026, 5, 11), None)
    agenda_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id="28718",
        file_id=None,
        attachment_id=None,
        source_url="https://example.test/agenda.pdf",
        content_hash="agenda-hash",
        processing_status=statuses.NEEDS_AGENDA_REVIEW,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        agenda_source,
        AgendaSegment(
            "28718",
            "100058",
            "91511",
            dt.date(2026, 5, 11),
            "Outdoor recreation to serve a 493-unit multi-family dwelling. ## Secretary's Report ## Upcoming Matters",
        ),
        AgendaClassification("100058", statuses.NEEDS_AGENDA_REVIEW, 0, "Boilerplate tail", "Secretary's Report"),
    )
    app_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="28718",
        file_id="91511",
        attachment_id="179990",
        source_url="https://example.test/application.pdf",
        content_hash="app-hash",
        processing_status=statuses.NEEDS_OPERATOR_REVIEW,
    )
    extraction_id = store.upsert_application_extraction(
        run_id,
        app_source,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url="https://example.test/application.pdf",
            attachment_id="179990",
            applicant=ContactFields(
                name="Joey Wisniewski",
                company="New Land Enterprises",
                mailing_address="1840A N. Farwell Ave, Milwaukee, WI 53202",
            ),
            project_contact=ContactFields(),
            owner=ContactFields(),
            section5_description="493-unit multi-family dwelling.",
            unit_count=493,
            status=statuses.NEEDS_OPERATOR_REVIEW,
            target_project=True,
        ),
    )

    assert "Agenda item is not currently classified as a hit" in store.list_application_extractions()[0]["quality_issues"]

    reviewed = store.review_agenda_item(agenda_id, statuses.AGENDA_HIT)
    app_row = store.list_application_extractions(statuses.APPLICATION_EXTRACTED)[0]

    assert reviewed["classification"] == statuses.AGENDA_HIT
    assert "Secretary" not in reviewed["description"]
    assert app_row["id"] == extraction_id
    assert app_row["quality_issues"] == []
    store.review_application(extraction_id, statuses.ACCEPTED, {}, None)
    assert store.list_application_extractions(statuses.ACCEPTED)[0]["id"] == extraction_id


def test_review_acceptance_rejects_uncertain_or_unmailable_rows(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 1, 1), dt.date(2026, 1, 31), None)
    source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id="1",
        file_id=None,
        attachment_id=None,
        source_url="https://example.test/agenda.pdf",
        content_hash="agenda-hash",
        processing_status=statuses.AGENDA_HIT,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        source_id,
        AgendaSegment("1", "1001", "9001", dt.date(2026, 1, 1), "Construct apartments"),
        AgendaClassification("1001", statuses.AGENDA_HIT, 0.9, "Housing", "Construct apartments"),
    )
    app_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="1",
        file_id="9001",
        attachment_id="a",
        source_url="https://example.test/application.pdf",
        content_hash="app-hash",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    extraction_id = store.upsert_application_extraction(
        run_id,
        app_source,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url="https://example.test/application.pdf",
            attachment_id="a",
            applicant=ContactFields(name="Jane Applicant"),
            project_contact=ContactFields(),
            owner=ContactFields(),
            section5_description="Construct apartments.",
            unit_count=None,
            status=statuses.NEEDS_OPERATOR_REVIEW,
            target_project=None,
        ),
    )

    try:
        store.review_application(extraction_id, statuses.ACCEPTED, {}, None)
    except ValueError as exc:
        assert "Target project is not confirmed" in str(exc)
    else:
        raise AssertionError("Expected acceptance QC failure")

    store.review_application(
        extraction_id,
        statuses.ACCEPTED,
        {
            "target_project": True,
            "applicant_company": "Applicant LLC",
            "applicant_mailing_address": "123 Main Street, Madison, WI 53703",
        },
        None,
    )

    assert store.list_application_extractions(statuses.ACCEPTED)[0]["id"] == extraction_id


def test_saved_review_corrections_clear_qc_and_accept_later(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 1, 1), dt.date(2026, 1, 31), None)
    source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id="1",
        file_id=None,
        attachment_id=None,
        source_url="https://example.test/agenda.pdf",
        content_hash="agenda-hash",
        processing_status=statuses.AGENDA_HIT,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        source_id,
        AgendaSegment("1", "1002", "9002", dt.date(2026, 1, 2), "Construct apartments"),
        AgendaClassification("1002", statuses.AGENDA_HIT, 0.9, "Housing", "Construct apartments"),
    )
    app_source = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id="1",
        file_id="9002",
        attachment_id="b",
        source_url="https://example.test/application-b.pdf",
        content_hash="app-hash-b",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    extraction_id = store.upsert_application_extraction(
        run_id,
        app_source,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url="https://example.test/application-b.pdf",
            attachment_id="b",
            applicant=ContactFields(name="Applicant name Jane Raw"),
            project_contact=ContactFields(),
            owner=ContactFields(),
            section5_description="Construct apartments.",
            unit_count=None,
            status=statuses.NEEDS_OPERATOR_REVIEW,
            target_project=None,
        ),
    )

    store.review_application(
        extraction_id,
        statuses.NEEDS_OPERATOR_REVIEW,
        {
            "target_project": True,
            "applicant_name": "",
            "applicant_company": "Known Developer LLC",
            "applicant_mailing_address": "123 Main Street, Madison, WI 53703",
        },
        "Corrected by operator",
    )
    corrected = store.list_application_extractions(statuses.NEEDS_OPERATOR_REVIEW)[0]

    assert corrected["quality_issues"] == []
    assert corrected["applicant_name"] == ""
    assert corrected["applicant_company"] == "Known Developer LLC"
    assert corrected["notes"] == "Corrected by operator"

    store.review_application(extraction_id, statuses.ACCEPTED, {}, None)
    accepted = store.list_application_extractions(statuses.ACCEPTED)[0]

    assert accepted["id"] == extraction_id
    assert accepted["applicant_company"] == "Known Developer LLC"
    assert accepted["quality_issues"] == []


def test_review_rows_report_duplicate_accepted_contacts(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 1, 1), dt.date(2026, 2, 28), None)
    _accepted_extraction(
        store,
        run_id,
        meeting_date=dt.date(2026, 1, 10),
        city_item_id="1001",
        applicant=ContactFields(
            name="Shared Contact",
            company="Shared Housing LLC",
            mailing_address="123 Main Street, Madison, Wisconsin",
        ),
        project_contact=ContactFields(),
    )
    pending_id = _pending_extraction(
        store,
        run_id,
        meeting_date=dt.date(2026, 2, 10),
        city_item_id="1002",
        applicant=ContactFields(
            name="Shared Contact",
            company="Shared Housing LLC",
            mailing_address="123 Main Street, Madison, Wisconsin",
        ),
    )

    row = [item for item in store.list_application_extractions(statuses.APPLICATION_EXTRACTED) if item["id"] == pending_id][0]

    assert row["duplicate_contacts"]
    assert "already saved" in row["duplicate_contacts"][0]["message"]


def test_watchdog_marks_stale_application_docling_run_and_preserves_failure(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), None)

    assert store.heartbeat_run(
        run_id,
        statuses.APPLICATION_DOCLING,
        "docling",
        "agenda_item:240",
        "Extracting application_240_123.pdf with Docling",
    )
    with sqlite3.connect(store.db_path) as conn:
        conn.execute("UPDATE runs SET heartbeat_at = ? WHERE id = ?", ("2026-01-01T00:00:00Z", run_id))

    marked = RunWatchdog(store, stale_after_seconds=1).audit_once()
    row = store.get_run(run_id)
    late_completion = store.finish_run(run_id, statuses.COMPLETED)

    assert marked[0]["status"] == statuses.FAILED_APPLICATION_DOCLING
    assert row["status"] == statuses.FAILED_APPLICATION_DOCLING
    assert "heartbeat timed out" in row["last_error"]
    assert late_completion is False


def test_watchdog_kills_stale_live_worker(monkeypatch, tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), None)
    store.register_run_worker(run_id, 99999)
    store.log_event(run_id, "worker_spawn", "runner", None, "Spawned run worker PID 99999")
    assert store.heartbeat_run(run_id, statuses.APPLICATION_DOCLING, "docling", "agenda_item:240", "Extracting stuck PDF")
    with sqlite3.connect(store.db_path) as conn:
        conn.execute("UPDATE runs SET heartbeat_at = ? WHERE id = ?", ("2026-01-01T00:00:00Z", run_id))
    killed = []

    monkeypatch.setattr("plan_commission_workbench.storage._pid_alive", lambda _pid: True)
    monkeypatch.setattr(RunWatchdog, "_kill_process_tree", lambda _self, pid: killed.append(pid) or f"killed {pid}")

    marked = RunWatchdog(store, stale_after_seconds=1).audit_once()
    events = store.list_run_events(run_id)

    assert marked[0]["pid_alive"] is True
    assert marked[0]["worker_spawned"] is True
    assert killed == [99999]
    assert any(event["stage"] == "watchdog_worker_kill" and "killed 99999" in event["message"] for event in events)


def test_watchdog_does_not_kill_legacy_server_pid(monkeypatch, tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 6, 1), dt.date(2026, 6, 2), None)
    store.register_run_worker(run_id, 88888)
    assert store.heartbeat_run(run_id, statuses.APPLICATION_DOCLING, "docling", "agenda_item:240", "Legacy stuck PDF")
    with sqlite3.connect(store.db_path) as conn:
        conn.execute("UPDATE runs SET heartbeat_at = ? WHERE id = ?", ("2026-01-01T00:00:00Z", run_id))
    killed = []

    monkeypatch.setattr("plan_commission_workbench.storage._pid_alive", lambda _pid: True)
    monkeypatch.setattr(RunWatchdog, "_kill_process_tree", lambda _self, pid: killed.append(pid) or f"killed {pid}")

    marked = RunWatchdog(store, stale_after_seconds=1).audit_once()

    assert marked[0]["worker_spawned"] is False
    assert killed == []


def test_pid_alive_uses_windows_query_without_os_kill(monkeypatch) -> None:
    killed = []

    monkeypatch.setattr("plan_commission_workbench.storage.os.name", "nt")
    monkeypatch.setattr("plan_commission_workbench.storage._windows_pid_alive", lambda pid: pid == 12345)
    monkeypatch.setattr("plan_commission_workbench.storage.os.kill", lambda pid, signal: killed.append((pid, signal)))

    assert _pid_alive(12345) is True
    assert killed == []


def test_label_export_keeps_newest_duplicate_contact_and_older_new_contact(tmp_path) -> None:
    store = ReviewStore(tmp_path / "workbench.db")
    store.initialize()
    run_id = store.create_run(dt.date(2026, 1, 1), dt.date(2026, 2, 28), None)
    older_id = _accepted_extraction(
        store,
        run_id,
        meeting_date=dt.date(2026, 1, 10),
        city_item_id="1001",
        applicant=ContactFields(
            name="Shared Contact",
            company="Shared Housing LLC",
            mailing_address="123 Main Street, Madison, Wisconsin",
        ),
        project_contact=ContactFields(
            name="Useful Older Contact",
            company="Useful Development LLC",
            mailing_address="456 State Street, Madison, Wisconsin",
        ),
    )
    newer_id = _accepted_extraction(
        store,
        run_id,
        meeting_date=dt.date(2026, 2, 10),
        city_item_id="1002",
        applicant=ContactFields(
            name="Shared Contact",
            company="Shared Housing LLC",
            mailing_address="123 Main Street, Madison, Wisconsin",
        ),
        project_contact=ContactFields(),
    )

    result = ExportService(store).export(tmp_path / "labels.docx", statuses.ACCEPTED)
    body = _docx_text(tmp_path / "labels.docx")

    assert result["row_count"] == 2
    assert body.count("Shared Contact") == 1
    assert "Useful Older Contact" in body
    assert any(
        issue["extraction_id"] == older_id and "outdated duplicate" in issue["reason"]
        for issue in result["qc_issues"]
    )
    assert not any(issue["extraction_id"] == newer_id and "outdated duplicate" in issue["reason"] for issue in result["qc_issues"])


def _accepted_extraction(
    store: ReviewStore,
    run_id: int,
    *,
    meeting_date: dt.date,
    city_item_id: str,
    applicant: ContactFields,
    project_contact: ContactFields,
) -> int:
    """Purpose: seed one accepted application row for export tests."""

    source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id=f"event-{city_item_id}",
        file_id=None,
        attachment_id=None,
        source_url=f"https://example.test/agenda-{city_item_id}.pdf",
        content_hash=f"agenda-hash-{city_item_id}",
        processing_status=statuses.AGENDA_HIT,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        source_id,
        AgendaSegment(f"event-{city_item_id}", city_item_id, city_item_id, meeting_date, "Construct housing"),
        AgendaClassification(city_item_id, statuses.AGENDA_HIT, 0.9, "Housing", "Construct housing"),
    )
    app_source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id=f"event-{city_item_id}",
        file_id=city_item_id,
        attachment_id=f"attachment-{city_item_id}",
        source_url=f"https://example.test/application-{city_item_id}.pdf",
        content_hash=f"application-hash-{city_item_id}",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    extraction_id = store.upsert_application_extraction(
        run_id,
        app_source_id,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url=f"https://example.test/application-{city_item_id}.pdf",
            attachment_id=f"attachment-{city_item_id}",
            applicant=applicant,
            project_contact=project_contact,
            owner=ContactFields(),
            section5_description="Construct housing.",
            unit_count=10,
            status=statuses.APPLICATION_EXTRACTED,
            target_project=True,
        ),
    )
    store.review_application(extraction_id, statuses.ACCEPTED, {}, None)
    return extraction_id


def _pending_extraction(
    store: ReviewStore,
    run_id: int,
    *,
    meeting_date: dt.date,
    city_item_id: str,
    applicant: ContactFields,
) -> int:
    """Purpose: seed one review-ready row without accepting it."""

    source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="agenda",
        event_id=f"event-{city_item_id}",
        file_id=None,
        attachment_id=None,
        source_url=f"https://example.test/agenda-{city_item_id}.pdf",
        content_hash=f"agenda-hash-{city_item_id}",
        processing_status=statuses.AGENDA_HIT,
    )
    agenda_id = store.upsert_agenda_item(
        run_id,
        source_id,
        AgendaSegment(f"event-{city_item_id}", city_item_id, city_item_id, meeting_date, "Construct housing"),
        AgendaClassification(city_item_id, statuses.AGENDA_HIT, 0.9, "Housing", "Construct housing"),
    )
    app_source_id = store.upsert_source_item(
        run_id=run_id,
        source_kind="application",
        event_id=f"event-{city_item_id}",
        file_id=city_item_id,
        attachment_id=f"attachment-{city_item_id}",
        source_url=f"https://example.test/application-{city_item_id}.pdf",
        content_hash=f"application-hash-{city_item_id}",
        processing_status=statuses.APPLICATION_EXTRACTED,
    )
    return store.upsert_application_extraction(
        run_id,
        app_source_id,
        ApplicationExtraction(
            agenda_item_id=agenda_id,
            source_url=f"https://example.test/application-{city_item_id}.pdf",
            attachment_id=f"attachment-{city_item_id}",
            applicant=applicant,
            project_contact=ContactFields(),
            owner=ContactFields(),
            section5_description="Construct housing.",
            unit_count=10,
            status=statuses.APPLICATION_EXTRACTED,
            target_project=True,
        ),
    )


def _docx_text(path) -> str:
    """Purpose: read table text from generated mailing-label DOCX files."""

    from docx import Document

    document = Document(path)
    parts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)
